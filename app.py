import io
import json
import math
import random
from copy import deepcopy
from datetime import datetime

import pandas as pd
import streamlit as st
import streamlit.components.v1 as components
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ============================================================
# إعداد التطبيق
# ============================================================
st.set_page_config(
    page_title="نظام جدولة قسم اللغة العربية",
    page_icon="📚",
    layout="wide",
    initial_sidebar_state="expanded",
)

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Tajawal:wght@400;500;700;800&display=swap');

html, body, [class*="css"], .stApp {
    font-family: 'Tajawal', sans-serif !important;
}
.stApp { direction: rtl; }
section[data-testid="stSidebar"] { direction: rtl; }
div[data-testid="stMetric"] { direction: rtl; }
.block-container { padding-top: 1.5rem; padding-bottom: 2rem; }
.small-note { color:#64748b; font-size:0.88rem; }
.badge {
    display:inline-block; padding:5px 11px; border-radius:999px;
    background:#eef2ff; margin:2px; font-size:0.85rem;
}
.schedule-card {
    border-radius:12px; padding:9px; min-height:70px;
    box-shadow:0 2px 8px rgba(15,23,42,.08);
    border:1px solid rgba(0,0,0,.08);
}
</style>
""", unsafe_allow_html=True)

# ============================================================
# الثوابت
# ============================================================
DAY_PATTERNS = ["السبت / الاثنين", "الأحد / الثلاثاء"]
ACTUAL_DAYS = ["السبت", "الاثنين", "الأحد", "الثلاثاء"]

def expand_day_pattern(pattern):
    if pattern == "السبت / الاثنين":
        return ["السبت", "الاثنين"]
    elif pattern == "الأحد / الثلاثاء":
        return ["الأحد", "الثلاثاء"]
    else:
        return []

TIME_OPTIONS = [
    {"start": "08:15", "end": "09:05", "label": "08:15 - 09:05"},
    {"start": "08:15", "end": "09:30", "label": "08:15 - 09:30"},
    {"start": "09:45", "end": "10:35", "label": "09:45 - 10:35"},
    {"start": "09:45", "end": "11:00", "label": "09:45 - 11:00"},
    {"start": "11:15", "end": "12:05", "label": "11:15 - 12:05"},
    {"start": "11:15", "end": "12:30", "label": "11:15 - 12:30"},
    {"start": "12:40", "end": "13:30", "label": "12:40 - 13:30"},
    {"start": "12:40", "end": "13:55", "label": "12:40 - 13:55"},
    {"start": "14:15", "end": "15:30", "label": "14:15 - 15:30"},
    {"start": "20:10", "end": "21:00", "label": "20:10 - 21:00"},
    {"start": "21:10", "end": "22:00", "label": "21:10 - 22:00"},
    {"start": "22:10", "end": "23:00", "label": "22:10 - 23:00"},
]
TIME_LABELS = [x["label"] for x in TIME_OPTIONS]
TIME_MAP = {x["label"]: x for x in TIME_OPTIONS}

PALETTE = [
    "#DBEAFE", "#DCFCE7", "#FEF3C7", "#FCE7F3", "#EDE9FE",
    "#CFFAFE", "#FFEDD5", "#E0E7FF", "#F3E8FF", "#D1FAE5",
    "#FEE2E2", "#E2E8F0", "#CCFBF1", "#FAE8FF", "#ECFCCB",
]

# ============================================================
# الحالة
# ============================================================
def default_state():
    return {
        "department": "قسم اللغة العربية",
        "semester": "الفصل الدراسي الأول",
        "academic_year": "2026/2027",
        "doctors": {},
        "courses": {},
        "schedule": [],
        "history": [],
        "future": [],
        "next_course_id": 1,
        "next_event_id": 1,
    }

if "data" not in st.session_state:
    st.session_state.data = default_state()

D = st.session_state.data

# ============================================================
# أدوات مساعدة
# ============================================================
def snapshot():
    return deepcopy({
        "doctors": D["doctors"],
        "courses": D["courses"],
        "schedule": D["schedule"],
        "next_course_id": D["next_course_id"],
        "next_event_id": D["next_event_id"],
    })

def restore_snapshot(s):
    D["doctors"] = deepcopy(s["doctors"])
    D["courses"] = deepcopy(s["courses"])
    D["schedule"] = deepcopy(s["schedule"])
    D["next_course_id"] = s["next_course_id"]
    D["next_event_id"] = s["next_event_id"]

def push_history():
    D["history"].append(snapshot())
    if len(D["history"]) > 50:
        D["history"] = D["history"][-50:]
    D["future"] = []

def undo():
    if not D["history"]:
        return False
    D["future"].append(snapshot())
    restore_snapshot(D["history"].pop())
    return True

def redo():
    if not D["future"]:
        return False
    D["history"].append(snapshot())
    restore_snapshot(D["future"].pop())
    return True

def minutes(t):
    h, m = map(int, t.split(":"))
    return h * 60 + m

def overlaps(a_start, a_end, b_start, b_end):
    return minutes(a_start) < minutes(b_end) and minutes(b_start) < minutes(a_end)

def slot_dict(label):
    return TIME_MAP[label]

def doctor_color(name):
    if name in D["doctors"]:
        return D["doctors"][name].get("color", PALETTE[0])
    return PALETTE[0]

def event_course(event):
    return D["courses"].get(event["course_id"], {})

def patterns_share_day(pattern1, pattern2):
    days1 = set(expand_day_pattern(pattern1))
    days2 = set(expand_day_pattern(pattern2))
    return not days1.isdisjoint(days2)

def doctor_available(doctor, pattern, label):
    info = D["doctors"].get(doctor)
    if not info:
        return False
    if pattern not in info.get("available_patterns", DAY_PATTERNS):
        return False
    s = slot_dict(label)
    for blocked in info.get("blocked_slots", []):
        if blocked["day"] in expand_day_pattern(pattern):
            if overlaps(s["start"], s["end"], blocked["start"], blocked["end"]):
                return False
    br = info.get("break")
    if br and br != "لا يوجد":
        b = slot_dict(br)
        if overlaps(s["start"], s["end"], b["start"], b["end"]):
            return False
    if "max_daily" in info:
        max_daily = info["max_daily"]
        for day in expand_day_pattern(pattern):
            daily_count = sum(
                1 for e in D["schedule"]
                if e["doctor"] == doctor and day in expand_day_pattern(e["day"])
            )
            if daily_count >= max_daily:
                return False
    return True

def event_conflicts(candidate, ignore_id=None):
    problems = []
    cstart = slot_dict(candidate["time"])["start"]
    cend = slot_dict(candidate["time"])["end"]
    course = D["courses"].get(candidate["course_id"], {})
    doctor = candidate["doctor"]
    pattern = candidate["day"]

    if not doctor_available(doctor, pattern, candidate["time"]):
        problems.append("الدكتور غير متاح في هذا النمط/الوقت، أو تجاوز الحد الأقصى اليومي، أو المحاضرة تقع في وقت استراحته أو وقت محظور.")

    required_meetings = 2
    current_meetings = sum(
        1 for e in D["schedule"]
        if e["course_id"] == candidate["course_id"] and e["id"] != ignore_id
    )
    if current_meetings >= required_meetings:
        problems.append("المادة وصلت إلى العدد المطلوب من اللقاءات الأسبوعية (2).")

    for e in D["schedule"]:
        if e["id"] == ignore_id:
            continue
        if not patterns_share_day(pattern, e["day"]):
            continue
        estart = slot_dict(e["time"])["start"]
        eend = slot_dict(e["time"])["end"]

        if overlaps(cstart, cend, estart, eend):
            if e["doctor"] == doctor:
                problems.append(f"تعارض مع محاضرة أخرى للدكتور {doctor}.")
            if e["course_id"] == candidate["course_id"]:
                problems.append("تعارض: للمادة نفسها محاضرة أخرى في الوقت نفسه.")
            if e["course_id"] == candidate["course_id"] and e.get("section") == candidate.get("section"):
                problems.append("تعارض في الشعبة نفسها.")

    return list(dict.fromkeys(problems))

def add_event(course_id, doctor, pattern, time_label, section=None):
    course = D["courses"][course_id]
    ev = {
        "id": D["next_event_id"],
        "course_id": course_id,
        "course": course["name"],
        "doctor": doctor,
        "day": pattern,
        "time": time_label,
        "section": section or course.get("section", "1"),
        "color": doctor_color(doctor),
    }
    D["next_event_id"] += 1
    return ev

def course_label(course):
    sec = course.get("section", "1")
    return f'{course["name"]} — شعبة {sec}'

def schedule_quality():
    if not D["schedule"]:
        return 0
    score = 100
    conflicts = 0
    unavailable = 0

    for e in D["schedule"]:
        problems = event_conflicts(e, ignore_id=e["id"])
        conflicts += len(problems)
        if not doctor_available(e["doctor"], e["day"], e["time"]):
            unavailable += 1

    score -= conflicts * 12
    score -= unavailable * 10

    by_doc_day = {}
    for e in D["schedule"]:
        for day in expand_day_pattern(e["day"]):
            by_doc_day.setdefault((e["doctor"], day), 0)
            by_doc_day[(e["doctor"], day)] += 1
    for doctor in D["doctors"]:
        vals = [v for (d, _), v in by_doc_day.items() if d == doctor]
        if vals and max(vals) > 3:
            score -= (max(vals) - 3) * 2

    return max(0, min(100, round(score)))

def build_smart_schedule():
    candidates = list(D["courses"].keys())
    candidates.sort(key=lambda cid: 0 if D["courses"][cid].get("doctor") else 1)

    new_schedule = []
    old_schedule = D["schedule"]
    D["schedule"] = []

    for cid in candidates:
        c = D["courses"][cid]
        doctor = c.get("doctor")
        if not doctor or doctor not in D["doctors"]:
            continue

        best = None
        best_score = -10**9

        preferred = D["doctors"][doctor].get("preferred_period", "كلاهما")
        for pattern in DAY_PATTERNS:
            for label in TIME_LABELS:
                cand = {
                    "id": -1,
                    "course_id": cid,
                    "course": c["name"],
                    "doctor": doctor,
                    "day": pattern,
                    "time": label,
                    "section": c.get("section", "1"),
                    "color": doctor_color(doctor),
                }
                problems = event_conflicts(cand, ignore_id=-1)
                if problems:
                    continue

                sc = 0
                start = minutes(slot_dict(label)["start"])
                if preferred == "صباحي" and start >= 17 * 60:
                    sc -= 40
                if preferred == "مسائي" and start < 17 * 60:
                    sc -= 40

                same_pattern = sum(
                    1 for e in D["schedule"] if e["doctor"] == doctor and e["day"] == pattern
                )
                sc -= same_pattern * 7

                if any(e["course_id"] == cid and e["day"] == pattern for e in D["schedule"]):
                    sc -= 35

                doctor_events = [
                    e for e in D["schedule"] if e["doctor"] == doctor and e["day"] == pattern
                ]
                for e in doctor_events:
                    gap = abs(start - minutes(slot_dict(e["time"])["start"]))
                    if gap <= 20:
                        sc += 3
                    elif gap > 180:
                        sc -= 5

                sc += random.random() * 0.01
                if sc > best_score:
                    best_score = sc
                    best = cand

        if best:
            best["id"] = D["next_event_id"]
            D["next_event_id"] += 1
            D["schedule"].append(best)
        else:
            D["schedule"] = old_schedule
            return 0

    for cid in D["courses"]:
        if not any(e["course_id"] == cid for e in D["schedule"]):
            D["schedule"] = old_schedule
            return 0

    return len(D["schedule"])

# ============================================================
# التصدير
# ============================================================
def excel_export():
    wb = Workbook()
    ws = wb.active
    ws.title = "الجدول"
    ws.sheet_view.rightToLeft = True

    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(DAY_PATTERNS)+1)
    ws.cell(1,1).value = f'{D["department"]} — {D["semester"]} {D["academic_year"]}'
    ws.cell(1,1).font = Font(bold=True, size=16)
    ws.cell(1,1).alignment = Alignment(horizontal="center")

    headers = ["الوقت"] + DAY_PATTERNS
    for j, h in enumerate(headers, 1):
        c = ws.cell(3,j,h)
        c.font = Font(bold=True, color="FFFFFF")
        c.fill = PatternFill("solid", fgColor="1F4E78")
        c.alignment = Alignment(horizontal="center", vertical="center")

    for i, label in enumerate(TIME_LABELS, 4):
        ws.cell(i,1,label).font = Font(bold=True)
        ws.cell(i,1).alignment = Alignment(horizontal="center", vertical="center")
        for j, pattern in enumerate(DAY_PATTERNS, 2):
            events = [e for e in D["schedule"] if e["day"] == pattern and e["time"] == label]
            text = "\n\n".join(
                f'{e["course"]}\n{e["doctor"]} — شعبة {e.get("section","1")}'
                for e in events
            ) or "-"
            c = ws.cell(i,j,text)
            c.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            if events:
                hexcolor = doctor_color(events[0]["doctor"]).replace("#","")
                c.fill = PatternFill("solid", fgColor=hexcolor)

    for col in range(1, len(headers)+1):
        ws.column_dimensions[get_column_letter(col)].width = 30 if col > 1 else 20

    wd = wb.create_sheet("البيانات")
    wd.sheet_view.rightToLeft = True
    cols = ["الرقم","المادة","الشعبة","الدكتور","النمط","الوقت"]
    for j,h in enumerate(cols,1):
        wd.cell(1,j,h).font = Font(bold=True)
    for i,e in enumerate(D["schedule"],2):
        vals = [e["id"],e["course"],e.get("section","1"),e["doctor"],e["day"],e["time"]]
        for j,v in enumerate(vals,1):
            wd.cell(i,j,v)

    out = io.BytesIO()
    wb.save(out)
    return out.getvalue()

def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill.replace("#",""))
    tcPr.append(shd)

def word_export():
    doc = Document()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'{D["department"]}\n{D["semester"]} — {D["academic_year"]}')
    r.bold = True
    r.font.size = Pt(18)

    table = doc.add_table(rows=1, cols=len(DAY_PATTERNS)+1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["الوقت"] + DAY_PATTERNS
    for i,h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, "1F4E78")
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for rr in p.runs:
                rr.bold = True
                rr.font.color.rgb = RGBColor(255,255,255)

    for label in TIME_LABELS:
        cells = table.add_row().cells
        cells[0].text = label
        for p in cells[0].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for j, pattern in enumerate(DAY_PATTERNS,1):
            events = [e for e in D["schedule"] if e["day"] == pattern and e["time"] == label]
            cells[j].text = "\n\n".join(
                f'{e["course"]}\n{e["doctor"]} — شعبة {e.get("section","1")}' for e in events
            ) or "-"
            for p in cells[j].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if events:
                set_cell_shading(cells[j], doctor_color(events[0]["doctor"]))

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()

def csv_export():
    rows = []
    for e in D["schedule"]:
        rows.append({
            "الرقم": e["id"],
            "المادة": e["course"],
            "الشعبة": e.get("section","1"),
            "الدكتور": e["doctor"],
            "النمط": e["day"],
            "الوقت": e["time"],
        })
    return pd.DataFrame(rows).to_csv(index=False).encode("utf-8-sig")

def html_print_export():
    rows = []
    for label in TIME_LABELS:
        cells = []
        for pattern in DAY_PATTERNS:
            events = [e for e in D["schedule"] if e["day"] == pattern and e["time"] == label]
            if events:
                content = "".join(
                    f'<div class="card" style="background:{doctor_color(e["doctor"])}">'
                    f'<b>{e["course"]}</b><br>{e["doctor"]}<br>'
                    f'<small>شعبة {e.get("section","1")}</small></div>' for e in events
                )
            else:
                content = "-"
            cells.append(f"<td>{content}</td>")
        rows.append(f"<tr><th>{label}</th>{''.join(cells)}</tr>")

    html = f"""<!doctype html>
<html lang="ar" dir="rtl">
<head><meta charset="utf-8">
<title>{D["department"]}</title>
<style>
body{{font-family:Arial,Tahoma,sans-serif;margin:20px;direction:rtl}}
h1,h2{{text-align:center}}
table{{border-collapse:collapse;width:100%;table-layout:fixed}}
th,td{{border:1px solid #94a3b8;padding:7px;text-align:center;vertical-align:middle}}
thead th{{background:#1f4e78;color:white}}
.card{{border-radius:8px;padding:8px;margin:3px;font-size:13px}}
button{{padding:10px 18px;margin-bottom:12px}}
@media print{{button{{display:none}} @page{{size:A4 landscape;margin:8mm}}}}
</style></head>
<body>
<button onclick="window.print()">🖨️ طباعة / حفظ PDF</button>
<h1>{D["department"]}</h1>
<h2>{D["semester"]} — {D["academic_year"]}</h2>
<table><thead><tr><th>الوقت</th>{''.join(f"<th>{x}</th>" for x in DAY_PATTERNS)}</tr></thead>
<tbody>{''.join(rows)}</tbody></table>
</body></html>"""
    return html.encode("utf-8")

def doctor_time_table_export():
    """تصدير جدول شبكي (دكاترة × أوقات) كملف HTML"""
    doctors = list(D["doctors"].keys())
    times = TIME_LABELS
    
    # بناء الصفوف: لكل دكتور صف، ولكل وقت خلية
    rows = []
    for doc in doctors:
        row_cells = []
        for t in times:
            # العثور على محاضرات هذا الدكتور في هذا الوقت
            events = [e for e in D["schedule"] if e["doctor"] == doc and e["time"] == t]
            if events:
                cell_content = "<br>".join(
                    f'{e["course"]}<br><small>{e["day"]} - شعبة {e.get("section","1")}</small>'
                    for e in events
                )
                row_cells.append(f'<td style="background:{doctor_color(doc)}">{cell_content}</td>')
            else:
                row_cells.append("<td></td>")
        rows.append(f"<tr><th>{doc}</th>{''.join(row_cells)}</tr>")

    html = f"""<!doctype html>
<html lang="ar" dir="rtl">
<head><meta charset="utf-8"><title>جدول الدكاترة</title>
<style>
body{{font-family:Arial,Tahoma,sans-serif;direction:rtl;margin:20px}}
h2{{text-align:center}}
table{{border-collapse:collapse;width:100%;table-layout:fixed;font-size:12px}}
th,td{{border:1px solid #333;padding:6px;text-align:center;vertical-align:middle}}
th{{background:#1f4e78;color:white}}
td{{min-width:80px}}
</style></head>
<body>
<h2>جدول الدكاترة والأوقات</h2>
<table>
<thead><tr><th>الدكتور</th>{''.join(f"<th>{t}</th>" for t in times)}</tr></thead>
<tbody>{''.join(rows)}</tbody>
</table>
<br><button onclick='window.print()'>طباعة / PDF</button>
</body></html>"""
    return html.encode("utf-8")

def doctor_html_export(doctor_name):
    events = [e for e in D["schedule"] if e["doctor"] == doctor_name]
    if not events:
        return None

    rows = []
    for e in events:
        for actual_day in expand_day_pattern(e["day"]):
            rows.append({
                "اليوم": actual_day,
                "الوقت": e["time"],
                "المادة": e["course"],
                "الشعبة": e.get("section", "1"),
            })
    rows.sort(key=lambda x: (ACTUAL_DAYS.index(x["اليوم"]), TIME_LABELS.index(x["الوقت"])))

    html = f"""<!doctype html>
<html lang="ar" dir="rtl">
<head><meta charset="utf-8"><title>جدول {doctor_name}</title>
<style>
body{{font-family:Arial,Tahoma,sans-serif;direction:rtl;margin:20px}}
h2{{text-align:center}}
table{{border-collapse:collapse;width:100%}}
th,td{{border:1px solid #333;padding:8px;text-align:center}}
th{{background:#1f4e78;color:white}}
</style></head>
<body>
<h2>جدول الدكتور: {doctor_name}</h2>
<table>
<tr><th>اليوم</th><th>الوقت</th><th>المادة</th><th>الشعبة</th></tr>
"""
    for r in rows:
        html += f"<tr><td>{r['اليوم']}</td><td>{r['الوقت']}</td><td>{r['المادة']}</td><td>{r['الشعبة']}</td></tr>"
    html += "</table><br><button onclick='window.print()'>طباعة / PDF</button></body></html>"
    return html.encode("utf-8")

# ============================================================
# الواجهة التفاعلية (بدون حفظ تلقائي للسحب)
# ============================================================
def render_drag_drop():
    doctors = list(D["doctors"].keys())
    times = TIME_LABELS

    # تجهيز بيانات المحاضرات لكل خلية (doctor, time)
    grid_data = {}
    for e in D["schedule"]:
        key = (e["doctor"], e["time"])
        if key not in grid_data:
            grid_data[key] = []
        grid_data[key].append(e)

    payload = {
        "doctors": doctors,
        "times": times,
        "grid": {f"{doc}||{t}": [
            {"id": e["id"], "course": e["course"], "day": e["day"], "section": e.get("section","1"), "color": e["color"]}
            for e in grid_data.get((doc, t), [])
        ] for doc in doctors for t in times}
    }

    component = f"""
<!doctype html>
<html dir="rtl">
<head>
<meta charset="utf-8">
<style>
*{{box-sizing:border-box}}
body{{margin:0;font-family:Tahoma,Arial,sans-serif;background:#f8fafc;color:#0f172a}}
.wrap{{direction:rtl;overflow:auto;border:1px solid #cbd5e1;border-radius:14px;background:white}}
.grid{{display:grid;grid-template-columns:150px repeat({len(times)},minmax(120px,1fr));min-width:800px}}
.cell{{border-left:1px solid #dbe3ec;border-bottom:1px solid #dbe3ec;min-height:60px;padding:4px}}
.head{{background:#1e3a5f;color:white;min-height:40px;font-weight:bold;text-align:center;display:flex;align-items:center;justify-content:center;position:sticky;top:0;z-index:5}}
.doctor{{background:#f1f5f9;font-weight:bold;display:flex;align-items:center;justify-content:center;text-align:center}}
.card{{border-radius:8px;padding:4px;margin:2px 0;cursor:grab;box-shadow:0 2px 7px rgba(0,0,0,.12);border:1px solid rgba(0,0,0,.10);font-size:11px;text-align:center;user-select:none}}
.card:active{{cursor:grabbing;opacity:.7}}
.card b{{font-size:12px}}
.drop{{background:#eff6ff!important;outline:2px dashed #3b82f6;outline-offset:-3px}}
.hint{{font-size:12px;color:#64748b;padding:8px 2px}}
</style>
</head>
<body>
<div class="hint">💡 اسحب بطاقة المادة إلى خلية دكتور/وقت أخرى للتعديل المؤقت. استخدم النموذج أدناه للحفظ الفعلي.</div>
<div class="wrap"><div class="grid" id="grid"></div></div>
<script>
const DATA = {json.dumps(payload, ensure_ascii=False)};
const grid = document.getElementById("grid");

function esc(s) {{
 return String(s).replace(/[&<>"']/g,m=>({{"&":"&amp;","<":"&lt;",">":"&gt;",'"':"&quot;","'":"&#039;"}}[m]));
}}

grid.innerHTML = "";
let h = document.createElement("div");
h.className="cell head"; h.textContent="الدكتور"; grid.appendChild(h);

DATA.times.forEach(time=>{{
 let x=document.createElement("div"); x.className="cell head"; x.textContent=time; grid.appendChild(x);
}});

DATA.doctors.forEach(doc=>{{
 let d=document.createElement("div"); d.className="cell doctor"; d.textContent=doc; grid.appendChild(d);
 DATA.times.forEach(time=>{{
   let c=document.createElement("div"); c.className="cell dropzone";
   c.dataset.doctor=doc; c.dataset.time=time;
   c.addEventListener("dragover",ev=>{{ev.preventDefault();c.classList.add("drop")}});
   c.addEventListener("dragleave",()=>c.classList.remove("drop"));
   c.addEventListener("drop",ev=>{{
      ev.preventDefault(); c.classList.remove("drop");
      const id=Number(ev.dataTransfer.getData("text/plain"));
      const card = document.querySelector(`.card[data-id="${{id}}"]`);
      if (card && c !== card.parentElement) {{
        c.appendChild(card);
      }}
   }});
   grid.appendChild(c);
   
   // إضافة البطاقات الموجودة
   const key = doc + "||" + time;
   if (DATA.grid[key]) {{
     DATA.grid[key].forEach(e => {{
       const card=document.createElement("div");
       card.className="card"; card.draggable=true;
       card.dataset.id=e.id;
       card.style.background=e.color;
       card.innerHTML="<b>"+esc(e.course)+"</b><br><small>"+esc(e.day)+" - شعبة "+esc(e.section)+"</small>";
       card.addEventListener("dragstart",ev=>ev.dataTransfer.setData("text/plain",String(e.id)));
       c.appendChild(card);
     }});
   }}
 }});
}});
</script>
</body></html>
"""
    components.html(component, height=760, scrolling=True)
    st.caption("ملاحظة: السحب يعدّل العرض فقط. لحفظ التعديلات بشكل دائم استخدم أدوات التعديل أدناه.")

# ============================================================
# بيانات تجريبية (10 دكاترة × 4 مواد)
# ============================================================
def load_demo_data():
    # إنشاء قائمة بأسماء الدكاترة
    doctor_names = [
        "د. أحمد الشريف",
        "د. خالد العتيبي",
        "د. سارة المطيري",
        "د. محمد الهاشمي",
        "د. فاطمة الزهراء",
        "د. عبدالله الراشد",
        "د. نورة السبيعي",
        "د. يوسف الحربي",
        "د. ريم القحطاني",
        "د. طلال الغامدي",
    ]
    
    # إعداد الدكاترة
    D["doctors"] = {}
    for i, name in enumerate(doctor_names):
        D["doctors"][name] = {
            "available_patterns": DAY_PATTERNS,
            "break": "لا يوجد" if i % 3 != 0 else "12:40 - 13:30",
            "preferred_period": "صباحي" if i % 2 == 0 else "مسائي",
            "max_daily": 3,
            "color": PALETTE[i % len(PALETTE)],
            "blocked_slots": [],
        }
    
    # قائمة بأسماء المواد (40 مادة)
    course_names = [
        "النحو العربي 1", "النحو العربي 2", "الصرف 1", "الصرف 2",
        "البلاغة 1", "البلاغة 2", "العروض والقافية", "النقد الأدبي القديم",
        "النقد الأدبي الحديث", "الأدب الجاهلي", "الأدب الأموي", "الأدب العباسي",
        "الأدب الأندلسي", "الأدب الحديث", "تذوق أدبي", "مكتبة عربية",
        "تحقيق التراث", "مناهج النقد", "الشعر الجاهلي", "الشعر الأموي",
        "الشعر العباسي", "النثر العباسي", "النثر الحديث", "قضايا أدبية",
        "دراسات لغوية", "أصول النحو", "مدارس نحوية", "معاجم عربية",
        "علم الدلالة", "علم الأصوات", "تاريخ الأدب", "أدب الأطفال",
        "أدب الرحلة", "أدب السيرة", "قواعد اللغة", "تدريبات لغوية",
        "الكتابة العربية", "الإنشاء والتعبير", "فقه اللغة", "علم اللغة العام",
    ]
    
    # توزيع المواد: 4 مواد لكل دكتور
    D["courses"] = {}
    cid = 1
    for i, doc in enumerate(doctor_names):
        for j in range(4):
            course_name = course_names[i*4 + j]
            D["courses"][f"C{cid}"] = {
                "name": course_name,
                "code": f"ARB{cid}",
                "section": "1",
                "doctor": doc,
                "hours": 3,
            }
            cid += 1
    
    D["next_course_id"] = cid
    D["next_event_id"] = 1
    D["schedule"] = []

# ============================================================
# رأس التطبيق
# ============================================================
st.title("📚 نظام جدولة قسم اللغة العربية")
st.caption("بناء الجدول • التوزيع الذكي • السحب والإفلات • كشف التعارضات • التصدير")

c1,c2,c3,c4 = st.columns(4)
c1.metric("الدكاترة", len(D["doctors"]))
c2.metric("المواد", len(D["courses"]))
c3.metric("المحاضرات المجدولة", len(D["schedule"]))
c4.metric("جودة الجدول", f'{schedule_quality()}/100')

# ============================================================
# الشريط الجانبي
# ============================================================
with st.sidebar:
    st.header("⚙️ إعدادات المشروع")
    D["department"] = st.text_input("اسم القسم", D["department"])
    D["semester"] = st.text_input("الفصل الدراسي", D["semester"])
    D["academic_year"] = st.text_input("العام الجامعي", D["academic_year"])

    st.divider()
    st.subheader("💾 المشروع")
    b1,b2 = st.columns(2)
    if b1.button("↩️ تراجع", use_container_width=True):
        if undo(): st.rerun()
    if b2.button("↪️ إعادة", use_container_width=True):
        if redo(): st.rerun()

    project_bytes = json.dumps(
        {k:v for k,v in D.items() if k not in ("history","future")},
        ensure_ascii=False, indent=2
    ).encode("utf-8")
    st.download_button("💾 حفظ المشروع", project_bytes, "جدول_قسم_اللغة_العربية.json", "application/json", use_container_width=True)

    uploaded = st.file_uploader("📂 استعادة مشروع JSON", type=["json"])
    if uploaded is not None:
        try:
            obj = json.load(uploaded)
            push_history()
            for k in ["department","semester","academic_year","doctors","courses","schedule","next_course_id","next_event_id"]:
                if k in obj: D[k] = obj[k]
            st.success("تمت استعادة المشروع.")
            st.rerun()
        except Exception:
            st.error("ملف المشروع غير صالح.")

    st.divider()
    st.subheader("🧪 بيانات تجريبية")
    if st.button("📥 تحميل بيانات تجريبية", use_container_width=True):
        if D["doctors"] or D["courses"]:
            st.warning("سيتم استبدال البيانات الحالية. هل أنت متأكد؟")
            if st.button("نعم، استبدل البيانات"):
                push_history()
                load_demo_data()
                st.success("تم تحميل البيانات التجريبية.")
                st.rerun()
        else:
            push_history()
            load_demo_data()
            st.success("تم تحميل البيانات التجريبية.")
            st.rerun()

    st.divider()
    st.subheader("🧑‍🏫 إضافة أستاذ")
    with st.form("doctor_form"):
        name = st.text_input("اسم الدكتور")
        available_patterns = st.multiselect(
            "الأنماط المتاحة (أيام الدوام)",
            DAY_PATTERNS,
            default=DAY_PATTERNS[:],
        )
        br = st.selectbox("وقت الاستراحة", ["لا يوجد"] + TIME_LABELS)
        period = st.selectbox("الفترة المفضلة", ["كلاهما","صباحي","مسائي"])
        max_daily = st.number_input("الحد الأعلى للمحاضرات في اليوم الفعلي", 1, 8, 3)
        submitted = st.form_submit_button("➕ حفظ الأستاذ", use_container_width=True)
        if submitted:
            if not name.strip():
                st.error("أدخل اسم الدكتور.")
            else:
                push_history()
                old = D["doctors"].get(name.strip(), {})
                D["doctors"][name.strip()] = {
                    "available_patterns": available_patterns if available_patterns else DAY_PATTERNS,
                    "break": br,
                    "preferred_period": period,
                    "max_daily": int(max_daily),
                    "color": old.get("color", PALETTE[len(D["doctors"]) % len(PALETTE)]),
                    "blocked_slots": old.get("blocked_slots", []),
                }
                st.success("تم حفظ بيانات الأستاذ.")
                st.rerun()

    if D["doctors"]:
        st.caption("الأساتذة المسجلون")
        for doc in D["doctors"]:
            st.markdown(
                f'<span class="badge" style="background:{doctor_color(doc)}">{doc}</span>',
                unsafe_allow_html=True
            )

    st.divider()
    st.subheader("📚 إضافة مادة")
    with st.form("course_form"):
        cname = st.text_input("اسم المادة")
        ccode = st.text_input("رمز المادة (اختياري)")
        section = st.text_input("الشعبة", "1")
        doctor_options = ["— غير محدد —"] + list(D["doctors"].keys())
        assigned = st.selectbox("الدكتور المكلف", doctor_options)
        hours = st.number_input("الساعات المعتمدة", 1, 6, 3)
        if assigned != "— غير محدد —" and assigned in D["doctors"]:
            doctor_courses = [c["name"] for c in D["courses"].values() if c.get("doctor") == assigned]
            if doctor_courses:
                st.info(f"المواد الحالية للدكتور {assigned}: " + "، ".join(doctor_courses))
            else:
                st.info(f"لا توجد مواد مسندة للدكتور {assigned} حاليًا.")
        cs = st.form_submit_button("➕ حفظ المادة", use_container_width=True)
        if cs:
            if not cname.strip():
                st.error("أدخل اسم المادة.")
            else:
                push_history()
                cid = f"C{D['next_course_id']}"
                D["next_course_id"] += 1
                D["courses"][cid] = {
                    "name": cname.strip(),
                    "code": ccode.strip(),
                    "section": section.strip() or "1",
                    "doctor": None if assigned.startswith("—") else assigned,
                    "hours": int(hours),
                }
                st.success("تمت إضافة المادة.")
                st.rerun()

    st.divider()
    st.subheader("🧠 التوزيع")
    if st.button("🚀 توليد أفضل جدول متاح", type="primary", use_container_width=True):
        push_history()
        n = build_smart_schedule()
        if n > 0:
            st.success(f"تم توزيع {n} مادة.")
        else:
            st.error("تعذر توليد جدول كامل بدون تعارضات. حاول تعديل القيود.")
        st.rerun()

    if st.button("🧹 تفريغ الجدول فقط", use_container_width=True):
        push_history()
        D["schedule"] = []
        st.rerun()

    if st.button("🗑️ مسح المشروع بالكامل", use_container_width=True):
        if st.warning("هل أنت متأكد؟ سيتم حذف جميع البيانات نهائيًا."):
            st.session_state.data = default_state()
            st.rerun()

# ============================================================
# التبويبات
# ============================================================
tabs = st.tabs([
    "🗓️ الجدول التفاعلي",
    "🧑‍🏫 الأساتذة",
    "📚 المواد",
    "⚠️ فحص الجدول",
    "📊 التقارير والتصدير",
])

# ------------------------------------------------------------
# الجدول
# ------------------------------------------------------------
with tabs[0]:
    st.subheader("🗓️ الجدول الدراسي التفاعلي")
    st.write("العرض: الأوقات كأعمدة والدكاترة كصفوف. اسحب أي بطاقة إلى خلية دكتور/وقت أخرى (تعديل مؤقت). للحفظ الفعلي استخدم النموذج أدناه.")

    if not D["schedule"]:
        st.info("لا توجد محاضرات مجدولة. أضف الأساتذة والمواد ثم اضغط «توليد أفضل جدول متاح».")
    else:
        render_drag_drop()

        st.divider()
        st.subheader("✏️ تعديل محاضرة")
        labels = [
            f'#{e["id"]} — {e["course"]} — {e["doctor"]} — {e["day"]} — {e["time"]}'
            for e in D["schedule"]
        ]
        selected = st.selectbox("اختر محاضرة", labels)
        eid = int(selected.split("#")[1].split(" ")[0])
        ev = next(e for e in D["schedule"] if e["id"] == eid)

        col1,col2,col3 = st.columns(3)
        with col1:
            ndoc = st.selectbox("الدكتور", list(D["doctors"].keys()), index=list(D["doctors"]).index(ev["doctor"]))
        with col2:
            nday = st.selectbox("النمط", DAY_PATTERNS, index=DAY_PATTERNS.index(ev["day"]))
        with col3:
            ntime = st.selectbox("الوقت", TIME_LABELS, index=TIME_LABELS.index(ev["time"]))

        nsec = st.text_input("الشعبة", ev.get("section","1"))

        candidate = deepcopy(ev)
        candidate.update({"doctor":ndoc, "day":nday, "time":ntime, "section":nsec})
        problems = event_conflicts(candidate, ignore_id=eid)
        if problems:
            st.warning(" | ".join(problems))
        else:
            st.success("الموقع الجديد لا يحتوي على تعارض إلزامي.")

        col_btn1, col_btn2 = st.columns(2)
        with col_btn1:
            if st.button("💾 حفظ التعديل", type="primary", use_container_width=True):
                push_history()
                ev.update({"doctor":ndoc,"day":nday,"time":ntime,"section":nsec,"color":doctor_color(ndoc)})
                st.success("تم حفظ التعديل.")
                st.rerun()
        with col_btn2:
            if st.button("🗑️ حذف المحاضرة", use_container_width=True):
                push_history()
                D["schedule"] = [e for e in D["schedule"] if e["id"] != eid]
                st.success("تم حذف المحاضرة.")
                st.rerun()

        st.divider()
        st.subheader("➕ إضافة محاضرة يدويًا")
        with st.form("add_manual_event"):
            man_course = st.selectbox("المادة", [f"{cid}: {course_label(c)}" for cid,c in D["courses"].items()])
            man_doctor = st.selectbox("الدكتور", list(D["doctors"].keys()))
            man_day = st.selectbox("النمط", DAY_PATTERNS)
            man_time = st.selectbox("الوقت", TIME_LABELS)
            man_section = st.text_input("الشعبة", "1")
            submitted_manual = st.form_submit_button("إضافة المحاضرة")
            if submitted_manual:
                course_id = man_course.split(":")[0]
                cand = {
                    "id": -1,
                    "course_id": course_id,
                    "course": D["courses"][course_id]["name"],
                    "doctor": man_doctor,
                    "day": man_day,
                    "time": man_time,
                    "section": man_section,
                    "color": doctor_color(man_doctor),
                }
                problems = event_conflicts(cand, ignore_id=-1)
                if not problems:
                    push_history()
                    ev = add_event(course_id, man_doctor, man_day, man_time, man_section)
                    D["schedule"].append(ev)
                    st.success("تمت إضافة المحاضرة.")
                    st.rerun()
                else:
                    st.error("تعذر الإضافة: " + " | ".join(problems))

# ------------------------------------------------------------
# الأساتذة
# ------------------------------------------------------------
with tabs[1]:
    st.subheader("🧑‍🏫 إدارة الأساتذة")
    if not D["doctors"]:
        st.info("لم تتم إضافة أساتذة بعد.")
    else:
        for doc, info in list(D["doctors"].items()):
            with st.expander(f"👤 {doc}", expanded=False):
                a,b,c = st.columns(3)
                a.write("**الأنماط المتاحة:** " + "، ".join(info["available_patterns"]))
                b.write("**الاستراحة:** " + info["break"])
                c.write("**الفترة:** " + info["preferred_period"])
                st.write(f'الحد الأعلى اليومي: **{info["max_daily"]}**')
                assigned_courses = [
                    c["name"] for c in D["courses"].values() if c.get("doctor") == doc
                ]
                st.write("المواد: " + ("، ".join(assigned_courses) if assigned_courses else "لا توجد"))

                st.write("**الأوقات المحظورة:**")
                if info.get("blocked_slots"):
                    for bs in info["blocked_slots"]:
                        st.markdown(f"- {bs['day']} : {bs['start']} - {bs['end']}")
                else:
                    st.caption("لا توجد أوقات محظورة إضافية.")

                st.markdown("**تعديل البيانات:**")
                with st.form(f"edit_doctor_{doc}"):
                    new_patterns = st.multiselect(
                        "الأنماط المتاحة",
                        DAY_PATTERNS,
                        default=info["available_patterns"],
                    )
                    new_break = st.selectbox(
                        "وقت الاستراحة",
                        ["لا يوجد"] + TIME_LABELS,
                        index=(["لا يوجد"] + TIME_LABELS).index(info["break"]),
                    )
                    new_period = st.selectbox(
                        "الفترة المفضلة",
                        ["كلاهما","صباحي","مسائي"],
                        index=["كلاهما","صباحي","مسائي"].index(info["preferred_period"]),
                    )
                    new_max_daily = st.number_input("الحد الأعلى اليومي", 1, 8, info["max_daily"])
                    submit_edit = st.form_submit_button("حفظ التعديلات")
                    if submit_edit:
                        push_history()
                        info.update({
                            "available_patterns": new_patterns,
                            "break": new_break,
                            "preferred_period": new_period,
                            "max_daily": new_max_daily,
                        })
                        st.success("تم تحديث بيانات الأستاذ.")
                        st.rerun()

                st.markdown("**إضافة وقت محظور:**")
                with st.form(f"add_blocked_{doc}"):
                    block_day = st.selectbox("اليوم الفعلي", ACTUAL_DAYS)
                    block_start = st.selectbox("من", TIME_LABELS)
                    block_end = st.selectbox("إلى", TIME_LABELS, index=len(TIME_LABELS)-1)
                    add_block_btn = st.form_submit_button("إضافة")
                    if add_block_btn:
                        s = slot_dict(block_start)
                        e = slot_dict(block_end)
                        if minutes(s["start"]) >= minutes(e["end"]):
                            st.error("وقت البداية يجب أن يكون قبل وقت النهاية.")
                        else:
                            push_history()
                            info.setdefault("blocked_slots", []).append({
                                "day": block_day,
                                "start": s["start"],
                                "end": e["end"],
                            })
                            st.success("تمت إضافة الوقت المحظور.")
                            st.rerun()

                if st.button(f"حذف {doc}", key=f"del_doc_{doc}"):
                    push_history()
                    for course in D["courses"].values():
                        if course.get("doctor") == doc:
                            course["doctor"] = None
                    D["doctors"].pop(doc)
                    D["schedule"] = [e for e in D["schedule"] if e["doctor"] != doc]
                    st.rerun()

# ------------------------------------------------------------
# المواد
# ------------------------------------------------------------
with tabs[2]:
    st.subheader("📚 المواد والشعب")
    if not D["courses"]:
        st.info("لم تتم إضافة مواد بعد.")
    else:
        rows = []
        for cid,c in D["courses"].items():
            scheduled = sum(1 for e in D["schedule"] if e["course_id"] == cid)
            rows.append({
                "الرمز": c.get("code",""),
                "المادة": c["name"],
                "الشعبة": c.get("section","1"),
                "الدكتور": c.get("doctor") or "غير محدد",
                "اللقاءات المجدولة": scheduled,
                "الساعات": c.get("hours",3),
            })
        st.dataframe(pd.DataFrame(rows), use_container_width=True, hide_index=True)

        st.divider()
        for cid,c in list(D["courses"].items()):
            with st.expander(f'📖 {course_label(c)}'):
                with st.form(f"edit_course_{cid}"):
                    new_name = st.text_input("اسم المادة", c["name"])
                    new_code = st.text_input("رمز المادة", c.get("code",""))
                    new_section = st.text_input("الشعبة", c.get("section","1"))
                    doctor_options = ["— غير محدد —"] + list(D["doctors"].keys())
                    current_doctor = c.get("doctor") or "— غير محدد —"
                    new_doctor = st.selectbox(
                        "الدكتور",
                        doctor_options,
                        index=doctor_options.index(current_doctor) if current_doctor in doctor_options else 0,
                    )
                    new_hours = st.number_input("الساعات المعتمدة", 1, 6, c.get("hours",3))
                    submit_edit_course = st.form_submit_button("حفظ التعديلات")
                    if submit_edit_course:
                        if not new_name.strip():
                            st.error("اسم المادة مطلوب.")
                        else:
                            push_history()
                            c.update({
                                "name": new_name.strip(),
                                "code": new_code.strip(),
                                "section": new_section.strip() or "1",
                                "doctor": None if new_doctor == "— غير محدد —" else new_doctor,
                                "hours": int(new_hours),
                            })
                            for e in D["schedule"]:
                                if e["course_id"] == cid:
                                    e["course"] = new_name.strip()
                                    e["section"] = new_section.strip() or "1"
                                    e["doctor"] = c["doctor"]
                                    e["color"] = doctor_color(c["doctor"]) if c["doctor"] else PALETTE[0]
                            st.success("تم تحديث المادة.")
                            st.rerun()

                if st.button("حذف المادة", key=f"del_course_{cid}"):
                    push_history()
                    D["courses"].pop(cid)
                    D["schedule"] = [e for e in D["schedule"] if e["course_id"] != cid]
                    st.rerun()

# ------------------------------------------------------------
# الفحص
# ------------------------------------------------------------
with tabs[3]:
    st.subheader("⚠️ فحص الجدول")
    if not D["schedule"]:
        st.info("لا يوجد جدول لفحصه.")
    else:
        all_issues = []
        for e in D["schedule"]:
            for p in event_conflicts(e, ignore_id=e["id"]):
                all_issues.append((e, p))

        for cid, c in D["courses"].items():
            if not any(e["course_id"] == cid for e in D["schedule"]):
                all_issues.append((None, f"المادة {c['name']} (شعبة {c.get('section','1')}) غير مجدولة."))

        if not all_issues:
            st.success("🟢 ممتاز: لا توجد تعارضات إلزامية في الجدول الحالي.")
        else:
            st.error(f"يوجد {len(all_issues)} تنبيه/تعارض.")
            for e,p in all_issues:
                if e:
                    st.markdown(f'🔴 **{e["course"]} — {e["doctor"]} — {e["day"]} {e["time"]}:** {p}')
                else:
                    st.markdown(f'⚠️ {p}')

        st.divider()
        st.subheader("📈 جودة الجدول")
        q = schedule_quality()
        st.progress(q / 100)
        st.write(f"**التقييم الحالي: {q}/100**")

        stats = []
        for doc in D["doctors"]:
            events = [e for e in D["schedule"] if e["doctor"] == doc]
            by_day = {}
            for e in events:
                for day in expand_day_pattern(e["day"]):
                    by_day[day] = by_day.get(day, 0) + 1
            stats.append({
                "الدكتور": doc,
                "عدد المواد": len(events),
                "أعلى عدد في يوم فعلي": max(by_day.values()) if by_day else 0,
                "الأيام الفعلية المستخدمة": len(by_day),
            })
        st.dataframe(pd.DataFrame(stats), use_container_width=True, hide_index=True)

# ------------------------------------------------------------
# التقارير والتصدير
# ------------------------------------------------------------
with tabs[4]:
    st.subheader("📊 التقارير والتصدير")

    if D["schedule"]:
        a,b,c,d = st.columns(4)
        a.download_button(
            "📊 Excel",
            excel_export(),
            "جدول_قسم_اللغة_العربية.xlsx",
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True,
        )
        b.download_button(
            "📝 Word",
            word_export(),
            "جدول_قسم_اللغة_العربية.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )
        c.download_button(
            "📄 نسخة للطباعة",
            html_print_export(),
            "جدول_قسم_اللغة_العربية.html",
            "text/html",
            use_container_width=True,
        )
        d.download_button(
            "📁 CSV",
            csv_export(),
            "جدول_قسم_اللغة_العربية.csv",
            "text/csv",
            use_container_width=True,
        )

        st.divider()
        st.subheader("📋 تصدير جدول الدكاترة × الأوقات")
        if st.button("توليد ملف HTML لجدول الدكاترة"):
            html_bytes = doctor_time_table_export()
            st.download_button(
                "⬇️ تنزيل جدول الدكاترة HTML",
                html_bytes,
                "جدول_الدكاترة_والاوقات.html",
                "text/html",
                use_container_width=True,
            )

        st.divider()
        st.subheader("👨‍🏫 جداول الأساتذة")
        for doc in D["doctors"]:
            events = [e for e in D["schedule"] if e["doctor"] == doc]
            if events:
                col1, col2 = st.columns([3,1])
                with col1:
                    st.markdown(f"**{doc}** — {len(events)} مادة")
                    rows = []
                    for e in events:
                        for actual_day in expand_day_pattern(e["day"]):
                            rows.append({
                                "اليوم": actual_day,
                                "الوقت": e["time"],
                                "المادة": e["course"],
                                "الشعبة": e.get("section","1"),
                            })
                    rows.sort(key=lambda x: (ACTUAL_DAYS.index(x["اليوم"]), TIME_LABELS.index(x["الوقت"])))
                    st.dataframe(pd.DataFrame(rows), use_container_width=True, hide_index=True)
                with col2:
                    html_bytes = doctor_html_export(doc)
                    if html_bytes:
                        st.download_button(
                            "🖨️ تصدير جدول الدكتور",
                            html_bytes,
                            f"جدول_{doc}.html",
                            "text/html",
                            use_container_width=True,
                        )
    else:
        st.info("ولّد الجدول أولًا لتظهر خيارات التصدير.")

st.divider()
st.caption("📚 نظام جدولة قسم اللغة العربية — يعمل محليًا داخل Streamlit، مع حفظ واستعادة المشروع وتصدير الجدول.")
